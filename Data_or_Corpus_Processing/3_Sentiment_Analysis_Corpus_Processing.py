#!/usr/bin/env python
# coding: utf-8

# In[9]:


# load packages
import pandas as pd
import nltk
from nltk.tokenize import sent_tokenize
from docx import Document


# In[10]:


df = pd.read_csv("/Users/henryasiamah/Desktop/work/cleaned.csv")
df.head()


# In[11]:


df.at[3000, 'cleaned_corpus']


# In[12]:


# import nltk
# import pandas as pd
# from nltk.tokenize import sent_tokenize
# from docx import Document

# # Download NLTK resources
# nltk.download('punkt')

# # Load DataFrame
# df = pd.read_csv("/Users/henryasiamah/Desktop/work/cleaned.csv")

# # Define query words for sentence extraction (all lowercase for case-insensitive matching)
# query_words = ["refugees", "refugee", "migrants", "migration", "immigration", 
#                "immigrant", "migrant", "immigrants", "asylum seeker", "asylum seekers"]

# def extract_sentences_with_context(text, query_words):
#     if pd.isna(text):  # Handle NaN values
#         return []
    
#     sentences = sent_tokenize(text)
#     contexts = []
    
#     for i, sentence in enumerate(sentences):
#         # Check if any query word is in the current sentence
#         if any(word in sentence.lower() for word in query_words):
#             # Get context before (if available)
#             context_before = sentences[i-1] if i > 0 else ""
            
#             # Get context after (if available)
#             context_after = sentences[i+1] if i < len(sentences)-1 else ""
            
#             # Combine the context
#             full_context = f"{context_before} {sentence} {context_after}".strip()
#             contexts.append({
#                 'context_before': context_before,
#                 'query_sentence': sentence,
#                 'context_after': context_after,
#                 'full_context': full_context
#             })
    
#     return contexts

# # Initialize lists to store the results
# dates, headlines, extracted_contents = [], [], []
# context_before_list, context_after_list = [], []
# sources, countries_organizations = [], []

# # Process each row in the DataFrame
# for _, row in df.iterrows():
#     date = row['Date']
#     headline = row['Headline']
#     source = row['Source']
#     country_organization = row['Country/Organization']
#     cleaned_corpus = row['cleaned_corpus']
    
#     # Extract sentences with context
#     contexts = extract_sentences_with_context(cleaned_corpus, query_words)
    
#     # Append each extracted context along with other column values
#     for context in contexts:
#         dates.append(date)
#         headlines.append(headline)
#         sources.append(source)
#         countries_organizations.append(country_organization)
#         extracted_contents.append(context['query_sentence'])
#         context_before_list.append(context['context_before'])
#         context_after_list.append(context['context_after'])

# # Create a new DataFrame with the extracted information
# result_df = pd.DataFrame({
#     'Date': dates,
#     'Headline': headlines,
#     'Source': sources,
#     'Country/Organization': countries_organizations,
#     'Context_Before': context_before_list,
#     'Query_Sentence': extracted_contents,
#     'Context_After': context_after_list
# })

# # Save the results to an Excel file
# result_df.to_excel(r"C:\Users\pc\Desktop\work\combined_sent.xlsx", index=False)

# # Save the results to a CSV file
# result_df.to_csv(r"C:\Users\pc\Desktop\work\combined_sent.csv", index=False)

# # # Save the results to a Word document
# # doc = Document()
# # for row in result_df.itertuples():
# #     doc.add_paragraph(f"Date: {row.Date}")
# #     doc.add_paragraph(f"Headline: {row.Headline}")
# #     doc.add_paragraph(f"Source: {row.Source}")
# #     doc.add_paragraph(f"Country/Organization: {row.Country_Organization}")
# #     doc.add_paragraph("Context:")
# #     doc.add_paragraph(f"Before: {row.Context_Before}")
# #     doc.add_paragraph(f"Query Sentence: {row.Query_Sentence}")
# #     doc.add_paragraph(f"After: {row.Context_After}")
# #     doc.add_paragraph("\n")

# # doc.save(r"C:\Users\pc\Desktop\work\combined_sent.docx")
# print('Extraction and saving completed successfully!')


# ## Sentence Extraction

# In[13]:


import nltk
import pandas as pd
from nltk.tokenize import sent_tokenize
from docx import Document

# Download NLTK resources
nltk.download('punkt')

# Load DataFrame
df = pd.read_csv("/Users/henryasiamah/Desktop/work/cleaned.csv")

# Define query words for sentence extraction (all lowercase for case-insensitive matching)
query_words = ["refugees", "refugee", "migrants", "migration", "immigration", 
               "immigrant", "migrant", "immigrants", "asylum seeker", "asylum seekers"]

def extract_sentences_with_context(text, query_words):
    if pd.isna(text):  # Handle NaN values
        return []
    
    sentences = sent_tokenize(text)
    contexts = []
    
    for i, sentence in enumerate(sentences):
        # Check if any query word is in the current sentence
        if any(word in sentence.lower() for word in query_words):
            # Get context before (if available)
            context_before = sentences[i-1] if i > 0 else ""
            
            # Get context after (if available)
            context_after = sentences[i+1] if i < len(sentences)-1 else ""
            
            # Combine the context into a single text
            full_context = " ".join(filter(None, [context_before, sentence, context_after]))
            
            contexts.append({
                'context_before': context_before,
                'query_sentence': sentence,
                'context_after': context_after,
                'full_context': full_context
            })
    
    return contexts

# Initialize lists to store the results
dates, headlines, extracted_contents = [], [], []
context_before_list, context_after_list, full_context_list = [], [], []
sources, countries_organizations = [], []

# Process each row in the DataFrame
for _, row in df.iterrows():
    date = row['Date']
    headline = row['Headline']
    source = row['Source']
    country_organization = row['Country/Organization']
    cleaned_corpus = row['cleaned_corpus']
    
    # Extract sentences with context
    contexts = extract_sentences_with_context(cleaned_corpus, query_words)
    
    # Append each extracted context along with other column values
    for context in contexts:
        dates.append(date)
        headlines.append(headline)
        sources.append(source)
        countries_organizations.append(country_organization)
        extracted_contents.append(context['query_sentence'])
        context_before_list.append(context['context_before'])
        context_after_list.append(context['context_after'])
        full_context_list.append(context['full_context'])

# Create a new DataFrame with the extracted information
result_df = pd.DataFrame({
    'Date': dates,
    'Headline': headlines,
    'Source': sources,
    'Country/Organization': countries_organizations,
    'Context_Before': context_before_list,
    'Query_Sentence': extracted_contents,
    'Context_After': context_after_list,
    'Full_Context': full_context_list
})

# Save the results to an Excel file
result_df.to_excel(r"C:\Users\pc\Desktop\work\extracted_sentences_with_context.xlsx", index=False)

# Save the results to a CSV file
result_df.to_csv(r"C:\Users\pc\Desktop\work\extracted_sentences_with_context.csv", index=False)

# # Save the results to a Word document
# doc = Document()
# for row in result_df.itertuples():
#     doc.add_paragraph(f"Date: {row.Date}")
#     doc.add_paragraph(f"Headline: {row.Headline}")
#     doc.add_paragraph(f"Source: {row.Source}")
#     doc.add_paragraph(f"Country/Organization: {row.Country_Organization}")
#     doc.add_paragraph("Context:")
#     doc.add_paragraph(f"Before: {row.Context_Before}")
#     doc.add_paragraph(f"Query Sentence: {row.Query_Sentence}")
#     doc.add_paragraph(f"After: {row.Context_After}")
#     doc.add_paragraph(f"Full Context: {row.Full_Context}")
#     doc.add_paragraph("\n")

# doc.save(r"C:\Users\pc\Desktop\work\extracted_sentences_with_context.docx")
print('Extraction and saving completed successfully!')


# In[14]:


print(f"Number of rows in result_df: {len(result_df)}")


# In[16]:


senti_data = pd.read_csv("/Users/henryasiamah/Desktop/work/extracted_sentences_with_context.csv")
senti_data.head()


# In[18]:


senti_data.at[42023, 'Full_Context']


# In[19]:


senti_data['Full_Context'].duplicated().sum()


# In[20]:


senti_data.shape


# In[21]:


# remove duplicates
senti_data.drop_duplicates(subset=['Full_Context'], inplace=True)


# In[22]:


senti_data.shape


# In[23]:


senti_data.count()


# In[24]:


senti_data.to_csv("/Users/henryasiamah/Desktop/work/senti_data_new_combined.csv", index= False)


# In[ ]:




